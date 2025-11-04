from __future__ import annotations
from datetime import datetime, timezone

import gzip
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Tuple
import zipfile
from xml.sax.saxutils import escape

from .utils import dumps_json, ensure_directories

EXPORT_DIR = Path(__file__).resolve().parent.parent / "export"
ensure_directories(EXPORT_DIR)

CONTENT_TYPES_BASE = """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/settings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>
"""

CONTENT_TYPES_WITH_FOOTER = CONTENT_TYPES_BASE.replace(
    "</Types>",
    "  <Override PartName=\"/word/footer1.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml\"/>\n</Types>",
)

RELS_MAIN = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""

DOCUMENT_RELS_NO_FOOTER = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings" Target="settings.xml"/>
</Relationships>
"""

DOCUMENT_RELS_WITH_FOOTER = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings" Target="settings.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="footer1.xml"/>
</Relationships>
"""

STYLES_XML = """<?xml version="1.0" encoding="UTF-8"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:qFormat/>
  </w:style>
</w:styles>
"""

SETTINGS_XML = """<?xml version="1.0" encoding="UTF-8"?>
<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:compat/>
</w:settings>
"""

APP_XML = """<?xml version="1.0" encoding="UTF-8"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>Med Writer</Application>
</Properties>
"""

CORE_XML_TEMPLATE = """<?xml version="1.0" encoding="UTF-8"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>Med Writer</dc:title>
  <dc:creator>Med Writer</dc:creator>
  <cp:lastModifiedBy>Med Writer</cp:lastModifiedBy>
  <dcterms:created xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:modified>
</cp:coreProperties>
"""

DOCUMENT_TEMPLATE = """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    {paragraphs}
    {sect_pr}
  </w:body>
</w:document>
"""

FOOTER_TEMPLATE = """<?xml version="1.0" encoding="UTF-8"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:p><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>
</w:ftr>
"""


def _compose_stamp(config: Dict[str, Any]) -> str:
    if not config.get("habilitar"):
        return ""
    nome = config.get("nome", "").strip()
    crm = config.get("crm", "").strip()
    uf = config.get("uf", "").strip()
    especialidade = config.get("especialidade", "").strip()
    if not any([nome, crm, especialidade]):
        return ""
    partes = []
    if nome:
        partes.append(nome)
    if crm:
        crm_text = f"CRM {uf} {crm}" if uf else f"CRM {crm}"
        partes.append(crm_text.strip())
    if especialidade:
        partes.append(especialidade)
    return " | ".join(partes)


def build_docx(texto: str, stamp_config: Dict[str, Any] | None = None) -> bytes:
    lines = texto.splitlines() or [""]
    paragraphs = "".join(
        f"<w:p><w:r><w:t xml:space=\"preserve\">{escape(line)}</w:t></w:r></w:p>" for line in lines
    )
    stamp_text = _compose_stamp(stamp_config or {})
    if stamp_text:
        sect_pr = "<w:sectPr><w:footerReference w:type=\"default\" r:id=\"rId2\"/></w:sectPr>"
        footer_xml = FOOTER_TEMPLATE.format(text=escape(stamp_text))
        document_rels = DOCUMENT_RELS_WITH_FOOTER
        content_types = CONTENT_TYPES_WITH_FOOTER
    else:
        sect_pr = "<w:sectPr/>"
        footer_xml = None
        document_rels = DOCUMENT_RELS_NO_FOOTER
        content_types = CONTENT_TYPES_BASE
    document_xml = DOCUMENT_TEMPLATE.format(paragraphs=paragraphs, sect_pr=sect_pr)
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    files: Dict[str, bytes] = {
        "[Content_Types].xml": content_types.encode("utf-8"),
        "_rels/.rels": RELS_MAIN.encode("utf-8"),
        "word/document.xml": document_xml.encode("utf-8"),
        "word/_rels/document.xml.rels": document_rels.encode("utf-8"),
        "word/styles.xml": STYLES_XML.encode("utf-8"),
        "word/settings.xml": SETTINGS_XML.encode("utf-8"),
        "docProps/app.xml": APP_XML.encode("utf-8"),
        "docProps/core.xml": CORE_XML_TEMPLATE.format(timestamp=timestamp).encode("utf-8"),
    }
    if footer_xml:
        files["word/footer1.xml"] = footer_xml.encode("utf-8")

    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path, data in files.items():
            zf.writestr(path, data)
    return buffer.getvalue()


def apply_stamp(document: Any, config: Dict[str, Any]) -> None:  # pragma: no cover - opcional
    try:
        from docx.shared import Pt  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("python-docx não disponível") from exc
    if not hasattr(document, "sections"):
        raise RuntimeError("Documento inválido para aplicar carimbo")
    stamp_text = _compose_stamp(config)
    if not stamp_text:
        return
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    run = paragraph.add_run(stamp_text)
    run.font.size = Pt(float(config.get("tamanho", 9)))


def export_json(data: Dict[str, Any], compact: bool = False, compress: bool = False) -> bytes:
    payload = dumps_json(data, compact=compact)
    if not compress:
        return payload
    return gzip.compress(payload)


def create_zip_bundle(
    json_bytes: bytes,
    docx_bytes: bytes,
    metadata: Dict[str, Any],
    base_name: str,
) -> Tuple[bytes, Path]:
    ensure_directories(EXPORT_DIR)
    zip_buffer = BytesIO()
    timestamp = metadata.get("_meta", {}).get("gerado_em") or metadata.get("timestamp") or datetime.now(timezone.utc).isoformat(timespec="seconds")
    safe_name = base_name.replace(" ", "_").lower()
    file_name = f"{safe_name}-{timestamp.replace(':', '-')}.zip"
    target_path = EXPORT_DIR / file_name
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("soap.json", json_bytes)
        zf.writestr("documento.docx", docx_bytes)
        zf.writestr("session.json", dumps_json(metadata))
    with target_path.open("wb") as fp:
        fp.write(zip_buffer.getvalue())
    zip_buffer.seek(0)
    return zip_buffer.getvalue(), target_path
